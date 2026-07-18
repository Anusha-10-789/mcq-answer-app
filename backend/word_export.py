import io
import re
from typing import List

from docx import Document
from docx.shared import RGBColor

from models import MCQResult

BLACK = RGBColor(0, 0, 0)
OPTION_LINE_RE = re.compile(r"^([A-Za-z])[.):]\s*(.*)$")
BARE_LETTER_LINE_RE = re.compile(r"^[A-Za-z]$")


def _add_option_run(paragraph, letter: str, content: str) -> None:
    letter_run = paragraph.add_run(f"{letter.upper()})")
    letter_run.bold = True
    letter_run.font.color.rgb = BLACK
    if content:
        content_run = paragraph.add_run(content)
        content_run.font.color.rgb = BLACK


def _add_question_body(paragraph, question_text: str) -> None:
    lines = question_text.split("\n")
    rendered_any = False
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # A bare option letter on its own line, with its content on the next line.
        if BARE_LETTER_LINE_RE.match(stripped) and i + 1 < len(lines) and lines[i + 1].strip():
            if rendered_any:
                paragraph.add_run().add_break()
            _add_option_run(paragraph, stripped, lines[i + 1].strip())
            rendered_any = True
            i += 2
            continue

        if rendered_any:
            paragraph.add_run().add_break()

        match = OPTION_LINE_RE.match(stripped)
        if match:
            letter, content = match.groups()
            _add_option_run(paragraph, letter, content)
        else:
            run = paragraph.add_run(lines[i])
            run.font.color.rgb = BLACK

        rendered_any = True
        i += 1


def build_mcq_docx(results: List[MCQResult]) -> io.BytesIO:
    document = Document()

    for index, result in enumerate(results, start=1):
        topic_paragraph = document.add_heading(level=2)
        topic_run = topic_paragraph.add_run(f"Q.{index}) {result.topic}:")
        topic_run.bold = False
        topic_run.font.color.rgb = BLACK

        question_paragraph = document.add_paragraph()
        _add_question_body(question_paragraph, result.question)

        answer_paragraph = document.add_paragraph()
        answer_run = answer_paragraph.add_run(f"Answer: {result.correct_answer}")
        answer_run.font.color.rgb = BLACK

        explanation_label_run = document.add_paragraph("Explanation:").runs[0]
        explanation_label_run.font.color.rgb = BLACK

        explanation_run = document.add_paragraph(result.explanation).runs[0]
        explanation_run.font.color.rgb = BLACK

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
